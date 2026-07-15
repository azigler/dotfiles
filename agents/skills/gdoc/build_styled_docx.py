#!/usr/bin/env python3
"""build_styled_docx.py — render a THEME-styled .docx from the same block spec
that render_styled_blocks.mjs renders into a Google Doc tab. One content spec ->
two matching outputs (a shareable Word doc + an in-doc tab). Themes live in
themes/<name>.json; content declares blocks by role.

    python3 build_styled_docx.py spec.json out.docx [--pdf]

spec.json is the same file render_styled_blocks.mjs takes (docId/tabId are
ignored here). See that script's header for the block/role schema.
"""

import json
import os
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))


def load_theme(spec):
    if spec.get("themeDef"):
        return spec["themeDef"]
    name = spec.get("theme", "editorial")
    with open(os.path.join(HERE, "themes", f"{name}.json")) as f:
        return json.load(f)


def rgb(pal, c):
    h = pal.get(c, c)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def resolve_block(theme, blk):
    R = (
        theme.get("roles", {}).get(blk.get("role"), {})
        if blk.get("role")
        else {}
    )
    para = {
        "sa": blk.get("sa", R.get("sa", 0)),
        "sb": blk.get("sb", R.get("sb", 6)),
        "bullet": blk.get("bullet", R.get("bullet", False)),
        "rule": blk.get("rule", R.get("rule")),
    }
    if para["rule"]:
        return para, [{"t": ""}]
    base = {
        k: R.get(k)
        for k in ("size", "bold", "italic", "underline", "color", "upper")
    }

    def mk(src, own_role):
        RR = theme.get("roles", {}).get(own_role, {}) if own_role else {}
        s = {**base, **RR}
        for k in (
            "size",
            "bold",
            "italic",
            "underline",
            "color",
            "link",
            "upper",
        ):
            if k in src:
                s[k] = src[k]
        t = src.get("text", "")
        if s.get("upper"):
            t = t.upper()
        return {
            "t": t,
            "size": s.get("size"),
            "bold": s.get("bold"),
            "italic": s.get("italic"),
            "underline": s.get("underline"),
            "color": s.get("color"),
            "link": s.get("link"),
        }

    runs = (
        [mk(r, r.get("role")) for r in blk["runs"]]
        if blk.get("runs")
        else [mk(blk, None)]
    )
    return para, runs


def add_rule(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), hexcolor)
    b.append(bot)
    pPr.append(b)


def add_hyperlink(p, url, text, color_hex, size, font):
    r_id = p.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color_hex)
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    if size:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rPr.append(sz)
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hAnsi"), font)
    rPr.append(rf)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    p._p.append(link)


def build(spec, out_path):
    theme = load_theme(spec)
    pal = theme.get("palette", {})
    font = theme.get("font", "Calibri")
    ink = rgb(pal, "ink") if "ink" in pal else RGBColor(0x1A, 0x1A, 0x2E)

    doc = Document()
    s = doc.sections[0]
    s.page_height, s.page_width = Inches(11), Inches(8.5)
    s.top_margin = s.bottom_margin = Inches(0.9)
    s.left_margin = s.right_margin = Inches(1.0)
    nrm = doc.styles["Normal"]
    nrm.font.name = font
    nrm.font.size = Pt(11)
    nrm.font.color.rgb = ink

    for blk in spec["blocks"]:
        para, runs = resolve_block(theme, blk)
        if para["rule"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(para["sa"])
            p.paragraph_format.space_after = Pt(para["sb"])
            add_rule(p, pal.get(para["rule"], para["rule"]))
            continue
        p = doc.add_paragraph(style="List Bullet" if para["bullet"] else None)
        p.paragraph_format.space_before = Pt(para["sa"])
        p.paragraph_format.space_after = Pt(para["sb"])
        for r in runs:
            if r.get("link"):
                add_hyperlink(
                    p,
                    r["link"],
                    r["t"],
                    pal.get(r.get("color", "link"), "3B36C9"),
                    r.get("size"),
                    font,
                )
                continue
            run = p.add_run(r["t"])
            run.font.name = font
            if r.get("size"):
                run.font.size = Pt(r["size"])
            run.font.bold = bool(r.get("bold"))
            run.font.italic = bool(r.get("italic"))
            if r.get("underline"):
                run.font.underline = True
            if r.get("color"):
                run.font.color.rgb = rgb(pal, r["color"])
    doc.save(out_path)
    return out_path


def export_pdf(docx_path):
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    outdir = os.path.dirname(os.path.abspath(docx_path)) or "."
    subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            outdir,
            docx_path,
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    return docx_path.rsplit(".", 1)[0] + ".pdf"


if __name__ == "__main__":
    with open(sys.argv[1]) as f:
        spec = json.load(f)
    build(spec, sys.argv[2])
    print("DOCX:", sys.argv[2])
    if "--pdf" in sys.argv:
        pdf = export_pdf(sys.argv[2])
        if pdf:
            print("PDF:", pdf)
